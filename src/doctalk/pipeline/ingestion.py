import uuid
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from doctalk.embedding import EmbeddingManager
from doctalk.models import Chunk
from doctalk.storage import VectorStoreManager
from docx import Document as DocxDocument


class IngestionPipeline:
    """Simple pipeline for ingesting documents into the vector store."""

    def __init__(
        self,
        embedder: EmbeddingManager,
        vector_store: VectorStoreManager,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        """Initialize the ingestion pipeline.

        :param embedder: EmbeddingManager instance
        :param vector_store: VectorStoreManager instance
        :param chunk_size: Size of text chunks
        :param chunk_overlap: Overlap between chunks
        """
        self.embedder = embedder
        self.vector_store = vector_store
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    def ingest(self, file_path: str) -> None:
        """Ingest a single document file.

        :param file_path: Path to the document file
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f'File not found: {file_path}')

        documents = self._load_document(str(path))
        if not documents:
            return

        chunks = self.splitter.split_documents(documents)

        texts = [chunk.page_content for chunk in chunks]
        embeddings = self.embedder.embed_batch(texts)

        document_name = path.name
        document_id = str(uuid.uuid4())
        chunk_objects = [
            Chunk(
                text=chunk.page_content,
                embedding=embedding,
                document_name=document_name,
                document_id=document_id,
                chunk_index=i,
                metadata=chunk.metadata,
            )
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))
        ]

        self.vector_store.insert_chunks(chunk_objects)

    def ingest_batch(self, file_paths: list[str]) -> None:
        """Ingest multiple document files.

        :param file_paths: List of file paths
        """
        for file_path in file_paths:
            self.ingest(file_path)

    def _load_document(self, file_path: str) -> list[Document]:
        """Load a document using the appropriate LangChain loader.

        :param file_path: Path to the document
        :return: List of Document objects
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.pdf':
            loader = PyPDFLoader(file_path)
        elif suffix in ['.txt', '.md']:
            loader = TextLoader(file_path, encoding='utf-8')
        elif suffix == '.docx':
            doc = DocxDocument(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            return [Document(page_content=text, metadata={'source': file_path})]
        else:
            raise ValueError(f'Unsupported file type: {suffix}')

        return loader.load()
